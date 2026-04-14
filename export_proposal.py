import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_svsu_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            
        # Tables (Simplified handling)
        elif '|' in line and '-' not in line and 'Service' in line:
            # This is a header row, we'll try to find the full table
            continue # I'll just skip the simplistic table parser and do a better one if needed
            # For this simple export, I'll just write paragraphs for now or use a basic table builder
            
        # Bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line[re.search(r'\d+\.\s', line).end():], style='List Number')
            
        # Normal text
        else:
            # Basic bold handling
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Note: For SVSU proposal, I'll just write the full text professionally.
    # Actually, let's just use the HTML-to-Word trick as it handles tables PERFECTLY in Word.
    # The previous script was good, I just need to run it.

    doc.save(docx_path)

# Actually, I'll stick to the HTML-to-DOC trick because it preserves MD tables flawlessly when opened in Word.
# DOCX libraries struggle with complex tables in automatic conversion.

if __name__ == "__main__":
    import markdown
    md_file = r"C:\Users\USER\.gemini\antigravity\brain\ec78750b-dfb1-43a8-ac9b-da490c15cdfe\SVSU_Official_Proposal_1Lakh.md"
    output_doc = r"C:\Users\USER\Desktop\BOT-SVSU\SVSU_Official_Proposal_1Lakh.doc"
    
    with open(md_file, 'r', encoding='utf-8') as f:
        text = f.read()
    
    html_content = markdown.markdown(text, extensions=['extra', 'tables', 'fenced_code'])
    
    styled_html = f"""
    <html>
    <head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #333; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #2c3e50; text-align: center; }}
        h2 {{ color: #2980b9; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #95a5a6; padding: 10px; text-align: left; }}
        th {{ background-color: #ecf0f1; font-weight: bold; color: #2c3e50; }}
        tr:nth-child(even) {{ background-color: #fdfdfd; }}
    </style>
    </head>
    <body>{html_content}</body>
    </html>
    """
    
    with open(output_doc, 'w', encoding='utf-8') as f:
        f.write(styled_html)
    print(f"Exported to {output_doc}")
